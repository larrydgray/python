import docx
doc = docx.Document('trucker_resume.docx')
print(len(doc.paragraphs))
print(doc.paragraphs[0].text)
print(len(doc.paragraphs[1].runs))
print(doc.paragraphs[1].runs[0].text)
print(doc.paragraphs[1].runs[1].text)
print(doc.paragraphs[1].runs[2].text)
print(doc.paragraphs[1].runs[3].text)
